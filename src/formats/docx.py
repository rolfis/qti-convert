"""
Microsoft Word 2007 Format
"""

from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_BREAK
from htmldocx import HtmlToDocx
from logzero import logger
import re
import config

def write_file(data, outfile):
    doc = Document()
    doc = setup_a4(doc)
    doc = setup_metadata(doc)

    html_parser = HtmlToDocx()

    for assessment in data['assessment']:
        doc.add_heading(assessment['metadata']['title'], 0)

        # handle error from missing description
        description = assessment['metadata'].get('description') or ''
        logger.info("Writing assessment: " + assessment['metadata']['title'])
        logger.info("with description: " + description)
        html_parser.add_html_to_document(description, doc)

        for question in assessment['question']:
            if 'title' in question:
                doc.add_heading(question['title'], 1)
            if 'image' in question:
                for img in question['image']:
                    doc.add_picture(img['href'].replace("%20", " "), width=Mm(100))
            if 'text' in question and question['text'] != None:
                this_question_text = re.sub('</*tbody>', '', question['text']) # See https://github.com/pqzx/html2docx/issues/1
                html_parser.add_html_to_document(this_question_text, doc)
            if 'answer' in question:
                if question['question_type'] == "multiple_dropdowns_question":
                    for aindex, group in enumerate(question['answer']):
                        options = []
                        for option in group['options']:
                            if option['display']:
                                if 'text' in answer and option['text'] != None:
                                    options.append(option['text'])
                                else:
                                    options.append("---")
                        doc.add_paragraph(str(aindex+1) + ": " + ", ".join(map(str, options)))
                elif question['question_type'] == "matching_question":
                    table = doc.add_table(rows=1, cols=2)
                    for index, answer in enumerate(question['answer']):
                        cell_0 = table.cell(0, 0)
                        if 'image' in answer:
                            for img in answer['image']:
                                cell_0.add_picture(img['href'].replace("%20", " "), height=Mm(10))
                        if 'text' in answer and answer['text'] != None:
                            cell_0.text = cell_0.text + ("\n" if cell_0.text != "" else "") + answer['text']
                        if index == 0:
                            cell_1 = table.cell(0, 1)
                            for option in answer['options']:
                                if 'image' in option:
                                    for img in option['image']:
                                        cell_1.add_picture(img['href'].replace("%20", " "), height=Mm(10))
                                if 'text' in option and option['text'] != None:
                                    cell_1.text = cell_1.text + ("\n" if cell_1.text != "" else "") + option['text']
                elif question['question_type'] == "calculated_question":
                    if config.calculated_display_var_set_in_text:
                        doc.add_paragraph(config.blanks_replace_str * config.blanks_answer_n)
                    else:
                        for index, answer in enumerate(question['answer']):
                            if answer['display'] and 'text' in answer and answer['text'] != None:
                                html_parser.add_html_to_document("<p>" + str(index+1) + ". " + answer['text'] + ": " + config.blanks_replace_str * 20 + "</p>", doc)
                else:
                    for index, answer in enumerate(question['answer']):
                        if answer['display']:                            
                            if 'image' in answer:
                                for img in answer['image']:
                                    html_parser.add_html_to_document("<p>" + str(index+1) + ".</p>", doc)
                                    doc.add_picture(img['href'].replace("%20", " "), height=Mm(10))
                            if 'text' in answer and answer['text'] != None:
                                html_parser.add_html_to_document("<p>" + str(index+1) + ". </p>" + answer['text'], doc)
                        else:
                            doc.add_paragraph(config.blanks_replace_str * config.blanks_answer_n)

            doc.add_page_break()

    doc.save(outfile)

def setup_a4(document):
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(30)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)
    return document

def setup_metadata(document):
    properties = document.core_properties
    properties.author = "qti-converter"
    properties.title = "Quiz export from LMS"
    return document
